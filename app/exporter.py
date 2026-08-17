from __future__ import annotations

import io
import os
import re
import shutil
import subprocess
import tempfile
import zipfile
from datetime import datetime
from html import escape
from pathlib import Path
from typing import Any

from .image_fetch import ImageFetchError, image_bytes_for_document


BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
SOURCE_REF_RE = re.compile(r"\[(\d+)\]")


def _show_references(article: dict[str, Any]) -> bool:
    spec = (article.get("generationMeta") or {}).get("writingSpec") or {}
    return bool(spec.get("citations"))


def _article_title(article: dict[str, Any], fallback: str = "数据要素") -> str:
    return str(article.get("recommendedTitle") or (article.get("titleCandidates") or [fallback])[0] or fallback)


class ExportError(RuntimeError):
    pass


def export_article(record: dict[str, Any], fmt: str) -> tuple[bytes, str, str]:
    article = record.get("article") or {}
    sources = record.get("sources") or []
    query = str(record.get("query") or "数据要素")
    title = _article_title(article, query)
    safe_title = _safe_filename(title)
    fmt = str(fmt or "").lower()
    if fmt == "docx":
        return build_docx(article, sources), f"{safe_title}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    if fmt == "pdf":
        return build_pdf(article, sources), f"{safe_title}.pdf", "application/pdf"
    raise ExportError("unsupported export format")


def validate_export_bytes(data: bytes, fmt: str, *, expect_images: bool = False) -> None:
    """Structural export validation before a file is sent to the browser."""
    fmt = str(fmt or "").lower()
    if fmt == "docx":
        if not data.startswith(b"PK") or len(data) < 500:
            raise ExportError("Word 文件生成校验失败：文件结构不完整")
        try:
            with zipfile.ZipFile(io.BytesIO(data)) as archive:
                bad = archive.testzip()
                if bad:
                    raise ExportError(f"Word 文件生成校验失败：压缩包损坏（{bad}）")
                names = set(archive.namelist())
                required = {"[Content_Types].xml", "word/document.xml", "word/styles.xml"}
                if not required.issubset(names):
                    raise ExportError("Word 文件生成校验失败：缺少必要文档结构")
                document_xml = archive.read("word/document.xml")
                if b"<w:document" not in document_xml or b"<w:body" not in document_xml:
                    raise ExportError("Word 文件生成校验失败：正文 XML 不完整")
                if expect_images and not any(name.startswith("word/media/") for name in names):
                    raise ExportError("Word 文件生成校验失败：预期配图未嵌入文档")
        except zipfile.BadZipFile as exc:
            raise ExportError("Word 文件生成校验失败：DOCX 压缩结构损坏") from exc
        return
    if fmt == "pdf":
        if not data.startswith(b"%PDF") or len(data) < 4000 or b"%%EOF" not in data[-2048:]:
            raise ExportError("PDF 文件生成校验失败：文件结构不完整")
        try:
            from pypdf import PdfReader
        except ImportError:
            # Structural validation is still possible without pypdf.
            if data.count(b"/Type /Page") < 1 and b"/Pages" not in data:
                raise ExportError("PDF 文件生成校验失败：没有可识别页面结构")
            return
        try:
            reader = PdfReader(io.BytesIO(data), strict=False)
            if not reader.pages:
                raise ExportError("PDF 文件生成校验失败：没有可读取页面")
            for page in reader.pages:
                _ = page.mediabox.width, page.mediabox.height
        except ExportError:
            raise
        except Exception as exc:
            raise ExportError(f"PDF 文件生成校验失败：{str(exc)[:120]}") from exc
        return
    raise ExportError("unsupported export format")


def build_docx(article: dict[str, Any], sources: list[dict[str, Any]]) -> bytes:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Inches, Pt, RGBColor
    except ImportError:
        return _build_docx_fallback(article, sources)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.74)
    section.bottom_margin = Inches(0.70)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)

    styles = doc.styles
    normal = styles["Normal"]
    # Publishing defaults: Chinese body copy uses Songti; titles/headings use
    # Heiti. Word will use the named fonts when installed and a compatible CJK
    # fallback otherwise. East-Asian font properties are set explicitly so the
    # choice survives Office / WPS / LibreOffice round-trips.
    normal.font.name = "SimSun"
    normal.font.size = Pt(11.2)
    normal.font.color.rgb = RGBColor(28, 28, 28)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.paragraph_format.line_spacing = 1.68
    normal.paragraph_format.space_after = Pt(7.5)
    normal.paragraph_format.widow_control = True

    for style_name, size in [("Title", 24.5), ("Heading 1", 16.5), ("Heading 2", 15.8), ("Heading 3", 13.8)]:
        style = styles[style_name]
        style.font.name = "SimHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor(18, 18, 18)

    title = _article_title(article)
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(title)
    run.bold = True
    run.font.name = "SimHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

    deck = str(article.get("deck") or "").strip()
    if deck:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(15)
        run = p.add_run(deck)
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(96, 119, 138)
        _bottom_border(p, color="CFE4F2", size="10")

    cover = article.get("coverImage")
    if isinstance(cover, dict) and cover.get("url"):
        _docx_add_image(doc, cover, max_width_inches=6.55, max_height_inches=4.25)

    for block in article.get("blocks") or []:
        kind = block.get("type")
        if kind == "heading":
            level = max(1, min(int(block.get("level") or 2), 3))
            # The article's markdown usually starts from ##, so normalize it visually.
            style_name = "Heading 2" if level >= 2 else "Heading 1"
            p = doc.add_paragraph(style=style_name)
            p.paragraph_format.space_before = Pt(17)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            p.paragraph_format.widow_control = True
            p.add_run(_plain_text(str(block.get("text") or ""))).bold = True
        elif kind == "paragraph":
            p = doc.add_paragraph()
            # Standard Chinese editorial typography: first-line indent by two
            # full-width characters (2 * 11 pt at the default body size).
            p.paragraph_format.first_line_indent = Pt(22.4)
            p.paragraph_format.widow_control = True
            _add_docx_inline(p, str(block.get("text") or ""))
        elif kind == "bullets":
            for item in block.get("items") or []:
                p = doc.add_paragraph(style="List Bullet")
                _add_docx_inline(p, str(item))
        elif kind == "image" and block.get("url"):
            _docx_add_image(doc, block, max_width_inches=6.15, max_height_inches=5.65)

    if sources and _show_references(article):
        p = doc.add_paragraph(style="Heading 2")
        p.paragraph_format.space_before = Pt(24)
        p.add_run("参考来源").bold = True
        for idx, src in enumerate(sources, start=1):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.05)
            label = f"[{idx}] {src.get('title') or '来源'}"
            source_url = str(src.get("url") or "")
            if source_url.startswith(("http://", "https://")):
                _add_hyperlink(p, label, source_url, color="285D84")
            else:
                run = p.add_run(label)
                run.font.color.rgb = RGBColor.from_string("285D84")
            meta = " · ".join(x for x in [str(src.get("source") or ""), str(src.get("publishedAt") or "")[:10]] if x)
            if meta:
                r = p.add_run(f"\n{meta}")
                r.font.size = Pt(8.5)
                r.font.color.rgb = RGBColor(123, 145, 161)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(f"数治攻关 · {datetime.now().strftime('%Y-%m-%d')}")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(140, 161, 177)

    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()


def build_pdf(article: dict[str, Any], sources: list[dict[str, Any]]) -> bytes:
    """Export PDF from the same canonical DOCX layout when LibreOffice exists.

    This makes Word/PDF typography and pagination agree instead of maintaining two
    visibly different layout engines. ReportLab remains a dependency-free fallback
    for machines without LibreOffice. Set DEG_PDF_ENGINE=reportlab to force it.
    """
    if str(os.environ.get("DEG_PDF_ENGINE") or "").lower() != "reportlab":
        converted = _pdf_from_docx_if_available(article, sources)
        if converted:
            return converted
    return _build_pdf_reportlab(article, sources)


def _pdf_from_docx_if_available(article: dict[str, Any], sources: list[dict[str, Any]]) -> bytes | None:
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        return None
    try:
        docx_bytes = build_docx(article, sources)
        with tempfile.TemporaryDirectory(prefix="deg-pdf-") as tmp:
            root = Path(tmp)
            src = root / "article.docx"
            src.write_bytes(docx_bytes)
            env = dict(os.environ)
            home = root / "home"
            profile = root / "lo-profile"
            home.mkdir(parents=True, exist_ok=True)
            profile.mkdir(parents=True, exist_ok=True)
            env["HOME"] = str(home)
            proc = subprocess.run(
                [
                    soffice, "--headless",
                    f"-env:UserInstallation={profile.resolve().as_uri()}",
                    "--convert-to", "pdf", "--outdir", str(root), str(src),
                ],
                stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=35, check=False, env=env,
            )
            out = root / "article.pdf"
            if proc.returncode == 0 and out.exists():
                data = out.read_bytes()
                if data.startswith(b"%PDF") and b"%%EOF" in data[-4096:]:
                    return data
    except Exception:
        return None
    return None


def _build_pdf_reportlab(article: dict[str, Any], sources: list[dict[str, Any]]) -> bytes:
    try:
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_LEFT
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import mm
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont
        from reportlab.platypus import CondPageBreak, HRFlowable, Image as RLImage, KeepTogether, Paragraph, SimpleDocTemplate, Spacer
    except ImportError:
        return _build_pdf_fallback(article, sources)

    try:
        pdfmetrics.getFont("STSong-Light")
    except Exception:
        pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))

    output = io.BytesIO()
    doc = SimpleDocTemplate(
        output,
        pagesize=A4,
        leftMargin=21 * mm,
        rightMargin=21 * mm,
        topMargin=19 * mm,
        bottomMargin=18 * mm,
        title=_article_title(article),
        author="数治攻关",
    )
    base = getSampleStyleSheet()
    title_style = ParagraphStyle("CNTitle", parent=base["Title"], fontName="STSong-Light", fontSize=22.0, leading=29, textColor=colors.HexColor("#111111"), alignment=TA_LEFT, spaceAfter=8, wordWrap="CJK")
    deck_style = ParagraphStyle("CNDeck", parent=base["BodyText"], fontName="STSong-Light", fontSize=10.5, leading=19, textColor=colors.HexColor("#60778A"), spaceAfter=12, wordWrap="CJK")
    heading_style = ParagraphStyle("CNHeading", parent=base["Heading2"], fontName="STSong-Light", fontSize=15.2, leading=22.8, textColor=colors.HexColor("#111111"), spaceBefore=15, spaceAfter=7, keepWithNext=True, wordWrap="CJK")
    body_style = ParagraphStyle("CNBody", parent=base["BodyText"], fontName="STSong-Light", fontSize=11.2, leading=20.2, textColor=colors.HexColor("#1C1C1C"), firstLineIndent=22.4, spaceAfter=8.5, wordWrap="CJK")
    caption_style = ParagraphStyle("CNCaption", parent=base["BodyText"], fontName="STSong-Light", fontSize=7.6, leading=11, textColor=colors.HexColor("#8198AA"), spaceAfter=7, wordWrap="CJK")
    source_style = ParagraphStyle("CNSource", parent=body_style, fontSize=8.5, leading=14, textColor=colors.HexColor("#60778A"), leftIndent=4 * mm)

    story: list[Any] = []
    title = _article_title(article)
    title_image = _pdf_heiti_title_image(title, 168 * mm)
    if title_image is not None:
        story.append(title_image)
        story.append(Spacer(1, 2 * mm))
    else:
        story.append(Paragraph(_pdf_inline(title), title_style))
    deck = str(article.get("deck") or "").strip()
    if deck:
        story.append(Paragraph(_pdf_inline(deck), deck_style))
    story.append(HRFlowable(width="100%", thickness=0.6, color=colors.HexColor("#CFE4F2"), spaceBefore=2, spaceAfter=12))

    cover = article.get("coverImage")
    if isinstance(cover, dict) and cover.get("url"):
        story.extend(_pdf_image_flowables(cover, 168 * mm, caption_style))

    blocks = list(article.get("blocks") or [])
    block_index = 0
    while block_index < len(blocks):
        block = blocks[block_index]
        kind = block.get("type")
        # A heading immediately followed by an image is a single editorial unit.
        # KeepTogether prevents the heading from being stranded at the bottom of
        # one page while its illustration starts the next page.
        if kind == "heading" and block_index + 1 < len(blocks):
            next_block = blocks[block_index + 1]
            if next_block.get("type") == "image" and next_block.get("url"):
                # Reserve a realistic heading+image footprint before the group.
                # CondPageBreak is more deterministic here than relying only on
                # nested KeepTogether flowables.
                story.append(CondPageBreak(118 * mm))
                bundle = [Paragraph(_pdf_inline(str(block.get("text") or "")), heading_style)]
                bundle.extend(_pdf_image_flowables(next_block, 168 * mm, caption_style))
                story.append(KeepTogether(bundle))
                block_index += 2
                continue
        if kind == "heading":
            story.append(Paragraph(_pdf_inline(str(block.get("text") or "")), heading_style))
        elif kind == "paragraph":
            story.append(Paragraph(_pdf_inline(str(block.get("text") or "")), body_style))
        elif kind == "bullets":
            for item in block.get("items") or []:
                story.append(Paragraph("• " + _pdf_inline(str(item)), ParagraphStyle("CNBullet", parent=body_style, firstLineIndent=0, leftIndent=4 * mm)))
        elif kind == "image" and block.get("url"):
            story.extend(_pdf_image_flowables(block, 168 * mm, caption_style))
        block_index += 1

    if sources and _show_references(article):
        source_flowables: list[Any] = [
            Spacer(1, 7 * mm),
            HRFlowable(width="100%", thickness=0.6, color=colors.HexColor("#CFE4F2"), spaceAfter=8),
            Paragraph("参考来源", heading_style),
        ]
        source_paragraphs: list[Any] = []
        for idx, src in enumerate(sources, start=1):
            title_text = str(src.get("title") or "来源")
            url = str(src.get("url") or "")
            meta = " / ".join(x.replace("·", "/") for x in [str(src.get("source") or ""), str(src.get("publishedAt") or "")[:10]] if x)
            text = f"[{idx}] {escape(title_text)}"
            if url.startswith(("http://", "https://")):
                text = f'<link href="{escape(url, quote=True)}" color="#285D84">{text}</link>'
            if meta:
                text += f"<br/><font color='#8198AA'>{escape(meta)}</font>"
            source_paragraphs.append(Paragraph(text, source_style))
        # Typical research briefs only have a handful of references. Keep a short
        # reference section together so the final page is intentional instead of
        # leaving one orphaned source on an otherwise empty page. Larger source
        # lists remain splittable while keeping the heading with the first entries.
        if len(source_paragraphs) <= 6:
            story.append(KeepTogether(source_flowables + source_paragraphs))
        else:
            lead_count = min(2, len(source_paragraphs))
            story.append(KeepTogether(source_flowables + source_paragraphs[:lead_count]))
            story.extend(source_paragraphs[lead_count:])

    def footer(canvas, doc_obj):
        canvas.saveState()
        canvas.setFont("STSong-Light", 7.5)
        canvas.setFillColor(colors.HexColor("#8CA1B1"))
        canvas.drawString(21 * mm, 10 * mm, "数治攻关")
        canvas.drawRightString(A4[0] - 21 * mm, 10 * mm, f"{doc_obj.page}")
        canvas.restoreState()

    doc.build(story, onFirstPage=footer, onLaterPages=footer)
    return output.getvalue()



def _pdf_heiti_title_image(text: str, max_width: float) -> Any | None:
    """Render the main PDF title with an installed CJK sans/Heiti font.

    ReportLab's built-in Chinese CID fallback is extremely reliable for body
    text, but it is Songti-like. For the main title we prefer an installed
    Heiti/sans font and rasterize only that heading; no font file is bundled or
    distributed with the project. If the runtime has no suitable font, the
    caller safely falls back to selectable STSong text.
    """
    try:
        from PIL import Image, ImageDraw, ImageFont
        from reportlab.platypus import Image as RLImage
    except Exception:
        return None

    candidates = [
        Path("C:/Windows/Fonts/simhei.ttf"),
        Path("C:/Windows/Fonts/msyhbd.ttc"),
        Path("/System/Library/Fonts/PingFang.ttc"),
        Path("/System/Library/Fonts/STHeiti Medium.ttc"),
        Path("/usr/share/fonts/opentype/noto/NotoSansCJK-Bold.ttc"),
    ]
    font_path = next((path for path in candidates if path.exists()), None)
    if font_path is None:
        return None
    try:
        font = ImageFont.truetype(str(font_path), 68, index=0)
    except Exception:
        try:
            font = ImageFont.truetype(str(font_path), 68)
        except Exception:
            return None

    canvas_width = 1800
    horizontal_pad = 10
    max_text_width = canvas_width - horizontal_pad * 2
    measure = Image.new("RGB", (32, 32), "white")
    measure_draw = ImageDraw.Draw(measure)

    lines: list[str] = []
    current = ""
    for ch in str(text or "数据要素治理"):
        if ch == "\n":
            if current:
                lines.append(current)
                current = ""
            continue
        candidate = current + ch
        box = measure_draw.textbbox((0, 0), candidate, font=font)
        width = max(0, box[2] - box[0])
        if current and width > max_text_width:
            lines.append(current)
            current = ch
        else:
            current = candidate
    if current:
        lines.append(current)
    lines = lines or [str(text or "数据要素治理")]

    line_height = 98
    canvas_height = max(116, line_height * len(lines) + 8)
    image = Image.new("RGB", (canvas_width, canvas_height), "white")
    draw = ImageDraw.Draw(image)
    y = 0
    for line in lines:
        draw.text((horizontal_pad, y), line, font=font, fill=(17, 17, 17))
        y += line_height
    buffer = io.BytesIO()
    image.save(buffer, format="PNG", optimize=True)
    buffer.seek(0)
    height = max_width * canvas_height / canvas_width
    return RLImage(buffer, width=max_width, height=height)


def _xml_escape(text: str) -> str:
    return escape(str(text or ""), quote=False)


def _docx_picture_xml(rel_id: str, filename: str, ident: int) -> str:
    return f'''<w:p><w:pPr><w:jc w:val="center"/></w:pPr><w:r><w:drawing><wp:inline xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture"><wp:extent cx="6000000" cy="3375000"/><wp:docPr id="{ident}" name="Picture {ident}"/><a:graphic><a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture"><pic:pic><pic:nvPicPr><pic:cNvPr id="0" name="{escape(filename)}"/><pic:cNvPicPr/></pic:nvPicPr><pic:blipFill><a:blip r:embed="{rel_id}"/><a:stretch><a:fillRect/></a:stretch></pic:blipFill><pic:spPr><a:xfrm><a:off x="0" y="0"/><a:ext cx="6000000" cy="3375000"/></a:xfrm><a:prstGeom prst="rect"><a:avLst/></a:prstGeom></pic:spPr></pic:pic></a:graphicData></a:graphic></wp:inline></w:drawing></w:r></w:p>'''


def _build_docx_fallback(article: dict[str, Any], sources: list[dict[str, Any]]) -> bytes:
    # Create a valid DOCX without python-docx so export never asks for an install.
    from pathlib import Path
    media = []
    blocks = article.get("blocks") or []
    image_blocks = []
    if isinstance(article.get("coverImage"), dict) and article.get("coverImage", {}).get("url"):
        image_blocks.append(article["coverImage"])
    image_blocks.extend([b for b in blocks if b.get("type") == "image" and b.get("url")])
    for block in image_blocks:
        try:
            data = image_bytes_for_document(str(block.get("url")), label=str(block.get("description") or "文章配图"))
            ext = "jpg" if data.startswith(b"\xff\xd8\xff") else "png"
            media.append((f"image{len(media)+1}.{ext}", data))
        except Exception:
            media.append((None, None))

    media_index = 0
    def add_image_rel():
        nonlocal media_index
        while media_index < len(media) and media[media_index][0] is None:
            media_index += 1
        if media_index >= len(media):
            return None
        rid = f"rIdImg{media_index+1}"; name = media[media_index][0]; media_index += 1
        return rid, name

    def paragraph(text: str, style: str = "Normal", indent: bool = True, bold: bool = False) -> str:
        ind = '<w:ind w:firstLineChars="200"/>' if indent and style == "Normal" else ''
        if bold:
            rpr = '<w:rPr><w:rFonts w:eastAsia="黑体"/><w:b/></w:rPr>'
        else:
            rpr = '<w:rPr><w:rFonts w:eastAsia="宋体"/></w:rPr>'
        return f'<w:p><w:pPr><w:pStyle w:val="{style}"/>{ind}</w:pPr><w:r>{rpr}<w:t xml:space="preserve">{_xml_escape(text)}</w:t></w:r></w:p>'

    body = []
    title = _article_title(article)
    body.append(paragraph(title, "Title", indent=False, bold=True))
    if article.get("deck"):
        body.append(paragraph(str(article.get("deck")), "Normal", indent=False))

    if isinstance(article.get("coverImage"), dict) and article.get("coverImage", {}).get("url"):
        rel = add_image_rel()
        if rel:
            rid, name = rel
            body.append(_docx_picture_xml(rid, name, 100))
            if article["coverImage"].get("caption") and str(article["coverImage"].get("provider") or "") != "generated-cover":
                body.append(paragraph(str(article["coverImage"]["caption"]), "Caption", indent=False))

    for block in blocks:
        kind = block.get("type")
        if kind == "heading":
            level = max(1, min(int(block.get("level") or 2), 3))
            body.append(paragraph(str(block.get("text") or ""), f"Heading{level}", indent=False, bold=True))
        elif kind == "paragraph":
            body.append(paragraph(str(block.get("text") or ""), "Normal", indent=True))
        elif kind == "bullets":
            for item in block.get("items") or []:
                body.append(paragraph("• " + str(item), "Normal", indent=False))
        elif kind == "image":
            rel = add_image_rel()
            if rel:
                rid, name = rel
                body.append(_docx_picture_xml(rid, name, 100+len(body)))
            if block.get("caption"):
                body.append(paragraph(str(block.get("caption")), "Caption", indent=False))

    if sources and _show_references(article):
        body.append(paragraph("参考来源", "Heading1", indent=False, bold=True))
        for i, src in enumerate(sources, 1):
            label = f"[{i}] {src.get('title') or '来源'}"
            if src.get("source"): label += f" · {src['source']}"
            if src.get("url"): label += f" · {src['url']}"
            body.append(paragraph(label, "Normal", indent=False))

    document_xml = '<?xml version="1.0" encoding="UTF-8" standalone="yes"?><w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"><w:body>' + ''.join(body) + '<w:sectPr><w:pgSz w:w="11906" w:h="16838"/><w:pgMar w:top="1020" w:right="1160" w:bottom="1020" w:left="1160"/></w:sectPr></w:body></w:document>'
    styles_xml = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?><w:styles xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:style w:type="paragraph" w:default="1" w:styleId="Normal"><w:name w:val="Normal"/><w:rPr><w:rFonts w:eastAsia="宋体"/><w:sz w:val="22"/></w:rPr></w:style><w:style w:type="paragraph" w:styleId="Title"><w:name w:val="Title"/><w:rPr><w:rFonts w:eastAsia="黑体"/><w:b/><w:sz w:val="48"/></w:rPr></w:style><w:style w:type="paragraph" w:styleId="Heading1"><w:name w:val="heading 1"/><w:rPr><w:rFonts w:eastAsia="黑体"/><w:b/><w:sz w:val="32"/></w:rPr></w:style><w:style w:type="paragraph" w:styleId="Heading2"><w:name w:val="heading 2"/><w:rPr><w:rFonts w:eastAsia="黑体"/><w:b/><w:sz w:val="30"/></w:rPr></w:style><w:style w:type="paragraph" w:styleId="Heading3"><w:name w:val="heading 3"/><w:rPr><w:rFonts w:eastAsia="黑体"/><w:b/><w:sz w:val="28"/></w:rPr></w:style><w:style w:type="paragraph" w:styleId="Caption"><w:name w:val="caption"/><w:rPr><w:rFonts w:eastAsia="宋体"/><w:sz w:val="16"/><w:color w:val="8198AA"/></w:rPr></w:style></w:styles>'''
    root_rels = '<?xml version="1.0" encoding="UTF-8" standalone="yes"?><Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"><Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/></Relationships>'
    rels = ['<?xml version="1.0" encoding="UTF-8" standalone="yes"?><Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"><Relationship Id="rIdStyles" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/styles" Target="styles.xml"/>']
    valid_media = [(i+1,m[0],m[1]) for i,m in enumerate(media) if m[0]]
    for idx,name,data in valid_media:
        rels.append(f'<Relationship Id="rIdImg{idx}" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/image" Target="media/{name}"/>')
    rels.append('</Relationships>')
    content_types = '<?xml version="1.0" encoding="UTF-8" standalone="yes"?><Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types"><Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/><Default Extension="xml" ContentType="application/xml"/><Default Extension="png" ContentType="image/png"/><Default Extension="jpg" ContentType="image/jpeg"/><Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/><Override PartName="/word/styles.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.styles+xml"/></Types>'
    out = io.BytesIO()
    with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as z:
        z.writestr('[Content_Types].xml', content_types); z.writestr('_rels/.rels', root_rels); z.writestr('word/document.xml', document_xml); z.writestr('word/styles.xml', styles_xml); z.writestr('word/_rels/document.xml.rels',''.join(rels))
        for _,name,data in valid_media:
            z.writestr(f'word/media/{name}', data)
    return out.getvalue()


def _build_pdf_fallback(article: dict[str, Any], sources: list[dict[str, Any]]) -> bytes:
    """Render a polished A4 PDF when ReportLab is unavailable.

    Earlier fallback output used one large sans-serif font for nearly everything,
    which made the PDF look unrelated to the in-browser preview. This renderer
    mirrors the editorial hierarchy instead: Heiti/sans for titles, Songti/serif
    for body copy, restrained spacing, first-line indentation and small captions.
    """
    try:
        from PIL import Image, ImageDraw, ImageFont
    except Exception as exc:
        raise ExportError("PDF 导出失败：当前环境缺少图像引擎") from exc

    sans_regular = _find_cjk_font(["NotoSansCJK-Regular.ttc", "msyh.ttc", "PingFang.ttc", "simhei.ttf"])
    sans_bold = _find_cjk_font(["NotoSansCJK-Bold.ttc", "msyhbd.ttc", "simhei.ttf", "PingFang.ttc"])
    serif_regular = _find_cjk_font(["NotoSerifCJK-Regular.ttc", "simsun.ttc", "simsun.ttf", "uming.ttc"])
    if not (sans_regular and serif_regular):
        raise ExportError("PDF 导出失败：没有可用中文字体")
    sans_bold = sans_bold or sans_regular

    W, H = 1240, 1754  # A4 at ~150dpi
    margin_x, top_y, bottom_y = 122, 100, 115
    content_w = W - margin_x * 2
    title_font = _load_cjk_font(sans_bold, 47)
    heading_font = _load_cjk_font(sans_bold, 31)
    body_font = _load_cjk_font(serif_regular, 23)
    deck_font = _load_cjk_font(serif_regular, 21)
    caption_font = _load_cjk_font(serif_regular, 16)
    footer_font = _load_cjk_font(sans_regular, 15)

    pages: list[Image.Image] = []
    page = Image.new("RGB", (W, H), "white")
    draw = ImageDraw.Draw(page)
    y = top_y
    page_no = 1

    def footer() -> None:
        draw.line((margin_x, H - 77, W - margin_x, H - 77), fill="#E7EEF6", width=1)
        draw.text((margin_x, H - 58), "数治攻关", font=footer_font, fill="#8CA1B1")
        number = str(page_no)
        tw = draw.textlength(number, font=footer_font)
        draw.text((W - margin_x - tw, H - 58), number, font=footer_font, fill="#8CA1B1")

    def new_page() -> None:
        nonlocal page, draw, y, page_no
        footer()
        pages.append(page)
        page_no += 1
        page = Image.new("RGB", (W, H), "white")
        draw = ImageDraw.Draw(page)
        y = top_y

    def ensure(height: int) -> None:
        if y + height > H - bottom_y:
            new_page()

    def wrap(text: str, font: Any, maxw: float, *, first_indent: float = 0.0) -> list[tuple[str, float]]:
        raw = _plain_text(str(text or "")).replace("\r", " ").replace("\n", " ")
        lines: list[tuple[str, float]] = []
        current = ""
        indent = first_indent
        for ch in raw:
            test = current + ch
            if current and draw.textlength(test, font=font) + indent > maxw:
                lines.append((current, indent))
                current = ch
                indent = 0.0
            else:
                current = test
        if current:
            lines.append((current, indent))
        return lines

    def add_text(text: str, font: Any, *, fill: str, line_h: int, first_indent: int = 0, gap_after: int = 0) -> None:
        nonlocal y
        lines = wrap(text, font, content_w, first_indent=first_indent)
        for line, indent in lines:
            ensure(line_h + 4)
            draw.text((margin_x + indent, y), line, font=font, fill=fill)
            y += line_h
        y += gap_after

    # Title + deck align to the browser preview rather than the old oversized fallback.
    title_lines = wrap(_article_title(article), title_font, content_w)
    for line, _ in title_lines:
        ensure(63)
        draw.text((margin_x, y), line, font=title_font, fill="#111827")
        y += 63
    y += 10
    deck = str(article.get("deck") or "").strip()
    if deck:
        add_text(deck, deck_font, fill="#60778A", line_h=35, gap_after=16)
    draw.line((margin_x, y, W - margin_x, y), fill="#CFE4F2", width=2)
    y += 26

    def add_image(block: dict[str, Any], *, is_cover: bool = False) -> None:
        nonlocal y
        try:
            data = image_bytes_for_document(str(block.get("url") or ""), label=str(block.get("description") or "文章配图"))
            im = Image.open(io.BytesIO(data)).convert("RGB")
        except Exception:
            return
        max_h = 430 if is_cover else 515
        im.thumbnail((content_w, max_h), Image.Resampling.LANCZOS)
        needed = im.height + 70
        ensure(needed)
        x = margin_x + (content_w - im.width) // 2
        page.paste(im, (x, y))
        y += im.height + 10
        caption = "" if str(block.get("provider") or "") == "generated-cover" else str(block.get("caption") or block.get("description") or "").strip()
        if caption:
            cap_lines = wrap(caption, caption_font, content_w)
            for line, _ in cap_lines[:3]:
                ensure(25)
                tw = draw.textlength(line, font=caption_font)
                draw.text((margin_x + max(0, (content_w - tw) / 2), y), line, font=caption_font, fill="#8198AA")
                y += 25
        y += 19

    def estimate_image_height(block: dict[str, Any], *, is_cover: bool = False) -> int:
        try:
            data = image_bytes_for_document(str(block.get("url") or ""), label=str(block.get("description") or "文章配图"))
            im = Image.open(io.BytesIO(data)).convert("RGB")
            max_h = 430 if is_cover else 515
            im.thumbnail((content_w, max_h), Image.Resampling.LANCZOS)
            caption = str(block.get("caption") or block.get("description") or "").strip()
            cap_h = min(3, len(wrap(caption, caption_font, content_w))) * 25 if caption else 0
            return im.height + 10 + cap_h + 27
        except Exception:
            return 220

    cover = article.get("coverImage")
    if isinstance(cover, dict) and cover.get("url"):
        add_image(cover, is_cover=True)

    blocks = list(article.get("blocks") or [])
    for block_index, block in enumerate(blocks):
        kind = block.get("type")
        if kind == "heading":
            # Do not leave a heading stranded at the bottom of a raster fallback
            # page. Reserve enough room for at least the first meaningful unit
            # that follows it, mirroring Word/ReportLab's keep-with-next behavior.
            next_block = blocks[block_index + 1] if block_index + 1 < len(blocks) else {}
            next_kind = next_block.get("type") if isinstance(next_block, dict) else ""
            # A body image can be tall even when its source file is wide; keep a
            # generous page footprint so the heading and figure start together.
            # The visual cost of a little white space is far smaller than an
            # orphaned heading at the foot of the page.
            reserve = max(900, estimate_image_height(next_block) + 80) if next_kind == "image" else 155
            ensure(90 + reserve)
            y += 28
            for line, _ in wrap(str(block.get("text") or ""), heading_font, content_w):
                draw.text((margin_x, y), line, font=heading_font, fill="#111111")
                y += 47
            y += 5
        elif kind == "paragraph":
            add_text(str(block.get("text") or ""), body_font, fill="#1C1C1C", line_h=39, first_indent=46, gap_after=13)
        elif kind == "bullets":
            for item in block.get("items") or []:
                add_text("• " + str(item), body_font, fill="#1C1C1C", line_h=39, first_indent=0, gap_after=6)
            y += 8
        elif kind == "image" and block.get("url"):
            y += 8
            add_image(block)

    if sources and _show_references(article):
        ensure(150)
        y += 28
        draw.line((margin_x, y, W - margin_x, y), fill="#CFE4F2", width=2)
        y += 30
        draw.text((margin_x, y), "参考来源", font=heading_font, fill="#111111")
        y += 50
        for i, src in enumerate(sources, 1):
            text = f"[{i}] {src.get('title') or '来源'}"
            add_text(text, caption_font, fill="#60778A", line_h=27, gap_after=4)

    footer()
    pages.append(page)
    out = io.BytesIO()
    pages[0].save(out, format="PDF", save_all=True, append_images=pages[1:], resolution=150.0)
    return out.getvalue()


def _find_cjk_font(names: list[str]) -> Path | None:
    roots = [Path("/usr/share/fonts/opentype/noto"), Path("/usr/share/fonts/truetype/arphic"), Path("C:/Windows/Fonts"), Path("/System/Library/Fonts")]
    # A few distro fonts live one directory below the arphic root.
    recursive_roots = [Path("/usr/share/fonts")]
    for name in names:
        for root in roots:
            candidate = root / name
            if candidate.exists():
                return candidate
        for root in recursive_roots:
            if root.exists():
                hit = next(root.rglob(name), None)
                if hit:
                    return hit
    return None


def _load_cjk_font(path: Path, size: int):
    from PIL import ImageFont
    # Noto CJK TTC index 2 is Simplified Chinese on common Linux builds. If a
    # platform uses a different collection layout, Pillow's default index remains
    # a valid fallback.
    try:
        if path.suffix.lower() == ".ttc" and "Noto" in path.name:
            return ImageFont.truetype(str(path), size, index=2)
        return ImageFont.truetype(str(path), size)
    except Exception:
        return ImageFont.truetype(str(path), size)


def _docx_add_image(
    doc: Any, block: dict[str, Any], *, max_width_inches: float, max_height_inches: float = 5.65,
) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor

    label = str(block.get("description") or block.get("caption") or "文章配图")
    try:
        data = image_bytes_for_document(str(block.get("url") or ""), label=label)
    except Exception:
        return

    width_inches = float(max_width_inches)
    height_inches: float | None = None
    try:
        from PIL import Image
        with Image.open(io.BytesIO(data)) as im:
            width_px, height_px = im.size
        if width_px > 0 and height_px > 0:
            height_inches = width_inches * height_px / width_px
            if height_inches > max_height_inches:
                height_inches = max_height_inches
                width_inches = height_inches * width_px / height_px
    except Exception:
        height_inches = None

    caption = "" if str(block.get("provider") or "") == "generated-cover" else str(block.get("caption") or block.get("description") or "").strip()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2 if caption else 7)
    p.paragraph_format.keep_with_next = bool(caption)
    run = p.add_run()
    try:
        if height_inches is not None:
            run.add_picture(io.BytesIO(data), width=Inches(width_inches), height=Inches(height_inches))
        else:
            run.add_picture(io.BytesIO(data), width=Inches(width_inches))
    except Exception:
        return
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(8)
        cp.paragraph_format.keep_together = True
        cr = cp.add_run(caption)
        cr.font.size = Pt(8)
        cr.font.color.rgb = RGBColor(129, 152, 170)


def _pdf_image_flowables(block: dict[str, Any], max_width: float, caption_style: Any) -> list[Any]:
    from PIL import Image
    from reportlab.platypus import Image as RLImage, KeepTogether, Paragraph, Spacer
    from reportlab.lib.units import mm

    label = str(block.get("description") or block.get("caption") or "文章配图")
    try:
        data = image_bytes_for_document(str(block.get("url") or ""), label=label)
    except Exception:
        return []
    buffer = io.BytesIO(data)
    try:
        with Image.open(buffer) as im:
            width_px, height_px = im.size
    except Exception:
        width_px, height_px = 1400, 788
    buffer.seek(0)
    width = max_width
    height = width * height_px / max(1, width_px)
    max_height = 92 * mm
    if height > max_height:
        scale = max_height / height
        width *= scale
        height *= scale
    image = RLImage(buffer, width=width, height=height)
    caption = "" if str(block.get("provider") or "") == "generated-cover" else str(block.get("caption") or block.get("description") or "").strip()
    group: list[Any] = [Spacer(1, 3 * mm), image]
    if caption:
        group.append(Paragraph(_pdf_inline(caption), caption_style))
    group.append(Spacer(1, 2 * mm))
    return [KeepTogether(group)]


def _add_docx_inline(paragraph: Any, text: str) -> None:
    cursor = 0
    for match in BOLD_RE.finditer(text):
        if match.start() > cursor:
            paragraph.add_run(text[cursor:match.start()])
        run = paragraph.add_run(match.group(1))
        run.bold = True
        cursor = match.end()
    if cursor < len(text):
        paragraph.add_run(text[cursor:])


def _add_hyperlink(paragraph: Any, text: str, url: str, *, color: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    if not url.startswith(("http://", "https://")):
        paragraph.add_run(text)
        return
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    props = OxmlElement("w:rPr")
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    props.append(c)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    props.append(u)
    run.append(props)
    text_el = OxmlElement("w:t")
    text_el.text = text
    run.append(text_el)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _bottom_border(paragraph: Any, *, color: str, size: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def _pdf_inline(text: str) -> str:
    normalized = _plain_text(text).replace("·", " / ").replace("—", "-").replace("–", "-")
    escaped = escape(normalized, quote=False)
    return SOURCE_REF_RE.sub(r"<font color='#285D84'>[\1]</font>", escaped)


def _plain_text(text: str) -> str:
    return BOLD_RE.sub(lambda m: m.group(1), str(text or ""))


def _safe_filename(title: str) -> str:
    clean = re.sub(r"[\\/:*?\"<>|\r\n]+", "-", str(title or "文章")).strip(" .-")
    return (clean or "数据要素文章")[:72]
