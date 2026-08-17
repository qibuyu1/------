from __future__ import annotations

import hashlib
import io
import json
import re
from html.parser import HTMLParser
from pathlib import Path
from typing import Any


class FileIngestError(ValueError):
    pass


MAX_FILE_BYTES = 12 * 1024 * 1024
MAX_TEXT_CHARS = 60_000
SUPPORTED_EXTENSIONS = {".txt", ".md", ".csv", ".json", ".html", ".htm", ".docx", ".pdf"}


class _TextHTMLParser(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self.parts: list[str] = []

    def handle_data(self, data: str) -> None:
        value = data.strip()
        if value:
            self.parts.append(value)


def ingest_file(filename: str, data: bytes, content_type: str = "") -> dict[str, Any]:
    name = Path(filename or "上传资料").name[:180]
    if not data:
        raise FileIngestError("文件是空的")
    if len(data) > MAX_FILE_BYTES:
        raise FileIngestError("单个文件不能超过 12MB")
    ext = Path(name).suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise FileIngestError("支持 TXT / Markdown / CSV / JSON / HTML / DOCX / PDF")

    if ext == ".docx":
        text = _read_docx(data)
    elif ext == ".pdf":
        text = _read_pdf(data)
    elif ext in {".html", ".htm"}:
        parser = _TextHTMLParser()
        parser.feed(_decode_text(data))
        text = "\n".join(parser.parts)
    elif ext == ".json":
        raw = _decode_text(data)
        try:
            text = json.dumps(json.loads(raw), ensure_ascii=False, indent=2)
        except Exception:
            text = raw
    else:
        text = _decode_text(data)

    text = _clean_text(text)[:MAX_TEXT_CHARS]
    if len(text.strip()) < 20:
        raise FileIngestError("没有从文件中提取到足够的正文内容")
    digest = hashlib.sha256(data).hexdigest()[:16]
    return {
        "id": f"upload-{digest}",
        "type": "upload",
        "title": name,
        "url": "",
        "source": "用户上传",
        "publishedAt": None,
        "snippet": text[:900],
        "rawContent": text,
        "authors": [],
        "citations": None,
        "readCount": None,
        "openAccess": None,
        "relevance": 1.0,
        "authorityScore": None,
        "freshnessScore": None,
        "score": 100,
        "images": [],
        "origin": "upload",
        "selectedByUser": True,
        "sourceVerified": True,
        "sourceStatus": "local",
        "fileName": name,
        "contentType": content_type,
        "fileBytes": len(data),
    }


def _decode_text(data: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030", "utf-16"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            pass
    return data.decode("utf-8", errors="replace")


def _read_docx(data: bytes) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise FileIngestError("读取 DOCX 需要安装 python-docx") from exc
    try:
        doc = Document(io.BytesIO(data))
    except Exception as exc:
        raise FileIngestError("DOCX 文件无法读取或已经损坏") from exc
    parts: list[str] = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text.strip())
    for table in doc.tables:
        for row in table.rows:
            values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if values:
                parts.append(" | ".join(values))
    return "\n".join(parts)


def _read_pdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise FileIngestError("读取 PDF 需要安装 pypdf") from exc
    try:
        reader = PdfReader(io.BytesIO(data))
        pages = []
        for page in reader.pages[:120]:
            pages.append(page.extract_text() or "")
        return "\n\n".join(pages)
    except Exception as exc:
        raise FileIngestError("PDF 文件无法读取；扫描版 PDF 暂不做 OCR") from exc


def _clean_text(text: str) -> str:
    value = str(text or "").replace("\x00", " ").replace("\r", "\n")
    value = re.sub(r"[ \t]+", " ", value)
    value = re.sub(r"\n{3,}", "\n\n", value)
    return value.strip()
